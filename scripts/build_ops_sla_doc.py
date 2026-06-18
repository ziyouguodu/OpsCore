from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "deliverables/系统运维保障方案_SLA99.99.docx"
FONT = "Microsoft YaHei"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color="000000", size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="D0D7DE", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths, header_fill="F2F4F7"):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color="0B2545", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) <= 10 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], str(value), size=9.2, align=align)
    doc.add_paragraph()
    return table


def add_kv_table(doc, pairs):
    rows = [(k, v) for k, v in pairs]
    return add_table(doc, ["项目", "内容"], rows, [1.55, 4.82])


def add_bullets(doc, items, style="List Bullet"):
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(10.5)


def add_numbered(doc, items):
    add_bullets(doc, items, style="List Number")


def add_callout(doc, title, body, fill="F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.37])
    set_table_borders(table, color="D8DEE6")
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(31, 58, 95)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.name = FONT
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r2.font.size = Pt(10)
    doc.add_paragraph()


def set_style_font(style, size=None, color=None, bold=None):
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size:
        style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], 11, "111827")
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Title", 22, "0B2545", 0, 8),
        ("Subtitle", 11, "4B5563", 0, 12),
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        set_style_font(styles[name], size, color, name.startswith("Heading") or name == "Title")
        styles[name].paragraph_format.space_before = Pt(before)
        styles[name].paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        set_style_font(styles[name], 10.5, "111827")
        styles[name].paragraph_format.space_after = Pt(4)
        styles[name].paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = "系统运维保障方案 | SLA 99.99%"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = FONT
    header.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left = footer.add_run("内部运维保障文件 | 第 ")
    page_field = OxmlElement("w:fldSimple")
    page_field.set(qn("w:instr"), "PAGE")
    page_run = OxmlElement("w:r")
    page_text = OxmlElement("w:t")
    page_text.text = "1"
    page_run.append(page_text)
    page_field.append(page_run)
    footer._p.append(page_field)
    right = footer.add_run(" 页")
    for run in (left, right):
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(107, 114, 128)


def add_title_page(doc):
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("系统运维保障方案")
    s = doc.add_paragraph(style="Subtitle")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run("面向 99.99% SLA 的生产稳定性、监控告警、应急响应与灾备保障体系")
    doc.add_paragraph()
    add_kv_table(doc, [
        ("适用范围", "生产环境核心业务系统、基础设施、数据库、中间件、网络、安全及运维流程"),
        ("SLA 目标", "系统月度可用性不低于 99.99%，月度不可用时间不超过 4.38 分钟"),
        ("版本", "V1.1 草案"),
        ("编制日期", "2026 年 5 月 21 日"),
        ("维护部门", "运维团队 / 技术保障团队"),
    ])
    add_callout(
        doc,
        "核心原则",
        "99.99% SLA 的本质不是“故障后快速修复”，而是通过冗余架构、自动化切换、严格变更、端到端监控和演练机制，将单点故障、误变更和恢复不确定性压缩到最低。",
    )
    doc.add_page_break()


def add_contents_page(doc):
    doc.add_heading("目录", level=1)
    sections = [
        "1. 方案目标",
        "2. 适用范围",
        "3. SLA 指标体系",
        "4. 系统架构现状与部署拓扑",
        "4.2 技术栈与基础设施现状",
        "5. 高可用保障措施",
        "6. 监控告警体系",
        "6.1 LLM / ASR / TTS 专项告警",
        "6.2 IDC 与私有化虚拟化专项告警",
        "6.3 告警响应闭环机制",
        "7. 变更与发布管理",
        "8. 备份与灾备",
        "9. 安全保障",
        "10. 容量与性能保障",
        "11. 分级应急预案清单",
        "12. 巡检制度",
        "13. 运维组织与值班机制",
        "14. 故障复盘与持续改进",
        "15. 附录 A：SLA 统计口径",
        "16. 附录 B：上线前稳定性检查清单",
        "17. 附录 C：故障复盘模板",
    ]
    add_bullets(doc, sections)
    doc.add_page_break()


def add_topology(doc):
    doc.add_heading("4. 系统架构现状与部署拓扑", level=1)
    doc.add_paragraph("本节描述系统生产架构的推荐落地形态。结合当前技术栈，系统同时包含 LLM、ASR、TTS 等智能能力服务，以及 IDC 机房物理机、私有化虚拟化平台、数据库、中间件和运维观测平台。核心服务不得以单实例承载生产流量，所有关键组件必须明确故障切换、降级绕行和人工接管方式。")
    add_table(doc, ["层级", "建议组件", "高可用要求"], [
        ("接入层", "DNS、WAF、防 DDoS、四/七层负载均衡、专线或互联网出口", "支持多线路接入、健康检查、故障节点自动摘除"),
        ("网关层", "API Gateway、Ingress、统一鉴权、限流与熔断组件", "至少双实例部署，配置限流、鉴权、熔断、重试与路由回退"),
        ("智能能力层", "LLM 推理服务、ASR 服务、TTS 服务、模型网关、提示词与音频处理服务", "多实例部署，支持模型降级、队列削峰、超时熔断和备用模型切换"),
        ("业务应用层", "业务服务、定时任务、后台管理服务、任务调度服务", "无状态多副本，支持滚动发布、灰度发布、水平扩容和快速回滚"),
        ("数据与中间件层", "关系型数据库、Redis、消息队列、对象存储、向量库或检索服务", "主备或集群部署，关键数据跨节点复制，具备自动或半自动故障转移"),
        ("基础设施层", "IDC 物理机、私有化虚拟化平台、共享存储、网络交换设备、机房电力与制冷", "主机、存储、网络、电力链路具备冗余，虚拟化集群启用 HA 和资源预留"),
        ("观测与运维层", "监控、日志、链路追踪、拨测、告警平台、堡垒机、审计平台", "独立于业务集群部署，告警链路多通道冗余，生产操作全量留痕"),
    ], [1.0, 2.1, 3.27])
    doc.add_heading("4.1 推荐逻辑拓扑", level=2)
    topo = (
        "用户 / 第三方系统\n"
        "        │\n"
        "DNS / WAF / 防 DDoS / 负载均衡\n"
        "        │\n"
        "API 网关 / 统一鉴权 / 限流熔断\n"
        "        │\n"
        "业务应用集群 / 任务调度集群\n"
        "        │\n"
        "模型网关 / 推理编排服务\n"
        " ┌──────┼──────┐\n"
        "LLM 服务  ASR 服务  TTS 服务\n"
        "        │\n"
        "Redis / MQ / 向量检索 / 对象存储\n"
        "        │\n"
        "数据库主备或集群 / 备份恢复\n"
        "        │\n"
        "IDC 物理机 / 私有化虚拟化集群 / 共享存储 / 网络设备"
    )
    p = doc.add_paragraph()
    run = p.add_run(topo)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    add_callout(doc, "拓扑落地要求", "任何处在主链路上的组件都需要明确：是否存在单点、是否支持健康检查、是否支持自动或半自动切换、切换期间是否满足 RTO/RPO。")
    doc.add_heading("4.2 技术栈与基础设施现状", level=2)
    add_table(doc, ["类别", "涉及组件", "重点风险", "保障重点"], [
        ("LLM", "模型网关、推理服务、GPU/CPU 推理节点、提示词模板、向量检索", "模型超时、上下文过长、推理队列积压、GPU 显存耗尽、输出质量异常", "限流熔断、队列削峰、备用模型、Prompt 版本回退、质量抽检"),
        ("ASR", "音频接入、格式转换、VAD、ASR 推理服务、结果后处理", "音频无法识别、识别延迟升高、格式不兼容、噪声导致准确率下降", "音频链路拨测、转码失败告警、延迟分位监控、人工回退通道"),
        ("TTS", "文本规范化、TTS 推理服务、音频合成、音频缓存", "合成失败、首包延迟过高、音色模型异常、音频文件生成失败", "首包延迟监控、失败率告警、缓存兜底、备用音色或模板播报"),
        ("IDC 物理机", "计算节点、GPU 服务器、管理口、带外管理、电源、风扇、磁盘、RAID", "硬件故障、温度异常、磁盘损坏、网卡异常、GPU 掉卡", "硬件传感器监控、备件管理、带外访问、节点隔离和迁移"),
        ("私有化虚拟化", "虚拟化管理平台、宿主机集群、虚拟机、共享存储、虚拟网络", "宿主机故障、资源超分、HA 失败、存储延迟、虚拟交换异常", "集群 HA、资源预留、虚机反亲和、存储链路冗余、管理面备份"),
    ], [0.85, 1.55, 2.0, 1.97])
    doc.add_heading("4.3 IDC 与私有化平台部署原则", level=2)
    add_bullets(doc, [
        "核心业务虚拟机应跨宿主机、跨机柜或跨故障域部署，避免同一宿主机、同一存储链路或同一交换机成为集中风险。",
        "LLM/ASR/TTS 推理节点应按照计算资源类型分池管理，GPU 节点、CPU 节点、通用业务节点分别设置容量水位和调度策略。",
        "虚拟化集群需启用 HA，关键虚拟机设置反亲和策略，并保留足够 N+1 或 N+2 资源冗余。",
        "IDC 网络、电力、制冷、存储链路和带外管理纳入统一巡检，不把机房基础设施视为黑盒。",
    ])


def add_specialized_monitoring(doc):
    doc.add_heading("6.1 LLM / ASR / TTS 专项告警", level=2)
    add_table(doc, ["能力域", "告警指标", "建议阈值", "处置动作"], [
        ("LLM", "请求失败率、P95/P99 延迟、首 Token 延迟、推理队列长度、上下文超限率", "失败率连续 3 分钟高于 3%；P99 超过基线 2 倍；队列持续增长", "开启限流和排队；切换备用模型；缩短上下文；暂停低优先级任务；扩容推理节点"),
        ("LLM 质量", "空回复率、拒答率、格式错误率、工具调用失败率、人工质检异常率", "超过业务基线或连续异常", "回退 Prompt/模型版本；关闭异常工具调用；启用规则兜底；通知算法和业务确认影响"),
        ("ASR", "识别失败率、音频转码失败率、端到端识别延迟、空文本率、低置信度比例", "失败率连续 3 分钟高于 2%；低置信度比例突增", "切换备用 ASR 服务；启用重试；降级为人工录入或延后处理；排查音频格式和采样率"),
        ("TTS", "合成失败率、首包延迟、音频生成耗时、缓存命中率、音频文件不可访问率", "失败率连续 3 分钟高于 2%；首包延迟超过 SLA", "切换备用音色或模板播报；启用缓存音频；暂停非核心合成；排查文本规范化和存储链路"),
        ("模型资源", "GPU 利用率、显存使用率、GPU 温度、GPU ECC 错误、推理进程存活", "显存持续高于 90%；ECC 错误增加；进程异常退出", "摘除异常节点；迁移流量；重启推理进程；联系硬件维护并保留日志"),
    ], [0.85, 2.1, 1.55, 1.87])
    doc.add_heading("6.2 IDC 与私有化虚拟化专项告警", level=2)
    add_table(doc, ["对象", "告警指标", "建议阈值", "处置动作"], [
        ("物理服务器", "CPU、内存、磁盘、RAID、电源、风扇、温度、带外管理状态", "温度超阈值、磁盘预测失败、RAID 降级、电源单路异常", "迁移业务或虚机；下线节点；更换备件；确认机柜电力与制冷"),
        ("GPU 服务器", "GPU 掉卡、显存、温度、功耗、驱动错误、推理进程状态", "GPU 不可见、温度超限、驱动重置、显存泄漏", "摘除推理节点；重启驱动或节点；切换备用节点；升级为硬件故障"),
        ("虚拟化平台", "宿主机状态、HA 事件、资源超分、虚机迁移失败、管理面不可用", "宿主机离线、HA 失败、资源池剩余不足 N+1", "触发虚机迁移；限制新建虚机；恢复管理面；通知平台管理员"),
        ("共享存储", "IOPS、吞吐、延迟、存储池水位、链路状态、快照失败", "延迟超过基线 2 倍；水位超过 80%；链路降级", "暂停低优先级任务；扩容存储；切换链路；检查多路径和控制器"),
        ("机房网络", "交换机端口、链路错误包、丢包率、延迟、BGP/路由、专线状态", "核心链路丢包、端口 flap、路由异常", "切换链路；隔离异常端口；联系网络值班；执行流量绕行"),
        ("机房环境", "UPS、电池、配电、空调、温湿度、漏水、消防", "供电单路异常、温度持续升高、UPS 告警", "启动机房应急；迁移关键业务；通知 IDC 现场人员；评估关停非核心负载"),
    ], [1.05, 2.05, 1.65, 1.62])
    doc.add_heading("6.3 告警响应闭环机制", level=2)
    add_numbered(doc, [
        "告警触发后，由一线值班在规定时限内确认告警真实性、影响范围、故障等级和初步责任域。",
        "P1/P2 告警必须进入应急群，记录告警时间、确认时间、处置动作、恢复时间和升级路径。",
        "同一故障导致多条告警时，由值班人员合并为一个事件单，避免多人重复处理和信息分散。",
        "处置过程中每 10 分钟同步一次状态；若 15 分钟内无明确恢复路径，自动升级到技术负责人。",
        "告警恢复后必须补充根因、临时措施、长期措施、负责人和截止时间，未闭环项进入稳定性问题台账。",
    ])
    add_callout(doc, "告警质量要求", "告警规则应同时覆盖“症状指标”和“根因指标”。例如 ASR 失败率升高是症状，音频转码失败、推理队列积压、GPU 节点异常、存储不可写可能是根因。值班大屏应把两类指标放在同一视图中。")


def add_incident_playbooks(doc):
    doc.add_heading("11. 分级应急预案清单", level=1)
    doc.add_paragraph("应急预案按故障等级和影响范围触发。处理原则为先止血、再恢复、后定位，先保障核心业务链路，再处理非核心功能和体验问题。")
    add_table(doc, ["等级", "触发条件", "响应时限", "恢复目标", "升级机制"], [
        ("P1", "核心业务整体不可用、生产数据丢失风险、数据库主库异常、LLM/ASR/TTS 主链路不可用、虚拟化集群或核心网络大面积故障", "5 分钟", "RTO ≤ 15 分钟", "立即拉起应急群，通知技术负责人、业务负责人和基础设施负责人"),
        ("P2", "核心功能局部不可用、错误率显著升高、性能严重下降、消息大面积积压、推理队列持续积压、单故障域资源异常", "10 分钟", "RTO ≤ 30 分钟", "通知二线专家，必要时升级 P1"),
        ("P3", "非核心功能异常、单节点故障、容量水位异常、个别接口波动、单台物理机或单个推理节点异常", "30 分钟", "当日恢复或制定修复计划", "运维负责人跟进，纳入问题台账"),
        ("P4", "一般风险提示、巡检异常、优化建议、低风险告警", "工作时间", "按计划处理", "周会跟踪闭环"),
    ], [0.55, 2.55, 0.85, 1.15, 1.27])
    doc.add_heading("11.1 P1 核心链路不可用预案", level=2)
    add_numbered(doc, [
        "值班人员在 5 分钟内确认告警真实性，判断影响范围、开始时间和受影响业务。",
        "立即建立应急沟通群，指定应急负责人、记录员、运维处理人、开发处理人和业务接口人。",
        "优先执行止血动作：切流、扩容、降级、熔断、回滚最近变更或隔离异常节点。",
        "若数据库或中间件主节点异常，按预案执行主备切换，并校验数据复制状态。",
        "服务恢复后持续观察至少 30 分钟，确认错误率、延迟、业务成功率恢复到正常阈值。",
    ])
    doc.add_heading("11.2 数据库异常预案", level=2)
    add_bullets(doc, [
        "连接数打满：限制非核心连接、扩容连接池上限前先排查慢 SQL 和异常任务。",
        "主库故障：确认复制延迟和数据一致性后执行主备切换，应用侧刷新连接配置。",
        "慢查询激增：启用只读降级、暂停低优先级任务，必要时临时增加索引或调整执行计划。",
        "误操作风险：立即冻结写入或隔离表，基于备份和 binlog 执行点位恢复。",
    ])
    doc.add_heading("11.3 发布异常预案", level=2)
    add_bullets(doc, [
        "灰度阶段出现异常，立即停止扩大灰度范围，保留现场日志和指标。",
        "错误率超过阈值，执行版本回滚或流量切回上一稳定版本。",
        "配置变更异常，按配置中心历史版本回退并触发应用热加载或重启。",
        "数据库变更异常，优先使用兼容性脚本和回滚脚本，不直接破坏现场数据。",
    ])
    doc.add_heading("11.4 流量突增预案", level=2)
    add_bullets(doc, [
        "确认流量来源，区分真实业务高峰、爬虫、攻击或异常重试风暴。",
        "扩容接入层、应用层和缓存层，必要时临时提高限流阈值或开启排队机制。",
        "对非核心接口启用降级，暂停报表、批处理、同步等低优先级任务。",
        "与业务侧同步容量边界和用户影响，必要时发布用户提示。",
    ])
    doc.add_heading("11.5 安全事件预案", level=2)
    add_bullets(doc, [
        "出现异常登录、批量探测、越权访问或 WAF 高危拦截时，立即通知安全负责人。",
        "对可疑 IP、账号、Token 执行临时封禁或强制失效。",
        "保留访问日志、审计日志和主机证据，避免覆盖关键现场。",
        "确认入侵风险后启动隔离、漏洞修复、密钥轮换和影响面排查。",
    ])
    doc.add_heading("11.6 LLM 服务异常预案", level=2)
    add_bullets(doc, [
        "推理超时或失败率升高：立即确认是否由模型节点、GPU 资源、上下文长度、向量检索、模型网关或下游工具调用引起。",
        "推理队列积压：暂停低优先级请求，启用排队提示或异步处理，必要时扩容推理节点或切换轻量模型。",
        "输出质量异常：回退模型版本、Prompt 模板或工具调用配置，启用规则兜底话术，并通知业务确认影响范围。",
        "GPU 节点异常：从负载均衡或调度池摘除节点，保留驱动日志、推理进程日志和 GPU 事件，转入硬件或驱动排查。",
        "模型网关不可用：切换备用网关或直连备用推理服务，同时收敛流量，避免调用方重试风暴。",
    ])
    doc.add_heading("11.7 ASR / TTS 服务异常预案", level=2)
    add_bullets(doc, [
        "ASR 识别失败：检查音频上传、格式转换、采样率、VAD、ASR 推理和后处理链路，确认是否存在单环节失败。",
        "ASR 延迟升高：启用队列削峰，限制长音频任务，必要时降级为延后识别或人工处理。",
        "TTS 合成失败：检查文本规范化、音色模型、推理服务、音频文件写入和对象存储访问状态。",
        "TTS 首包延迟过高：优先启用缓存音频、模板播报或备用音色，降低用户等待感知。",
        "音频链路异常：保留样本音频、TraceId、模型版本和转码日志，便于算法和工程团队联合复盘。",
    ])
    doc.add_heading("11.8 IDC 物理机故障预案", level=2)
    add_bullets(doc, [
        "物理机硬件告警：确认电源、风扇、温度、磁盘、RAID、网卡和带外管理状态，判断是否需要立即下线。",
        "承载虚拟机的宿主机异常：优先迁移或重启关键虚拟机，禁止在原因不明时反复强制重启宿主机。",
        "GPU 服务器故障：先摘除推理服务流量，再进行驱动、进程、温度、显存和硬件状态排查。",
        "机房环境异常：同步 IDC 现场人员，评估电力、制冷、网络影响范围，必要时迁移关键业务并关闭非核心负载。",
        "硬件更换后：执行基础巡检、压力测试和业务探活，确认稳定后再纳入生产资源池。",
    ])
    doc.add_heading("11.9 私有化虚拟化平台故障预案", level=2)
    add_bullets(doc, [
        "虚拟化管理面不可用：确认业务虚拟机是否仍在运行，优先保障业务面，不盲目重启管理组件。",
        "宿主机离线或 HA 失败：核对受影响虚拟机清单，按业务优先级手动拉起或迁移关键虚拟机。",
        "共享存储延迟升高：暂停备份、快照、批处理等低优先级 IO，检查存储控制器、多路径和链路状态。",
        "虚拟网络异常：按租户、VLAN、虚拟交换机、物理上联链路逐层定位，必要时将核心服务切换到备用网络。",
        "资源池容量不足：冻结非必要扩容和新建虚拟机，释放低优先级资源，保障核心业务和模型推理资源。",
    ])
    doc.add_heading("11.10 告警平台自身故障预案", level=2)
    add_bullets(doc, [
        "监控平台不可用时，立即切换到备用告警通道和基础探活脚本，确保 P1/P2 事件仍可被发现。",
        "日志平台不可用时，保留本地日志滚动策略和关键节点登录路径，避免故障现场丢失。",
        "告警风暴时，先确认是否为真实大面积故障，再启用告警聚合、降噪和事件合并，保留核心症状告警。",
        "监控恢复后补采故障时间段关键指标，并在复盘中评估监控盲区和告警规则有效性。",
    ])


def build():
    doc = Document()
    configure_doc(doc)
    add_title_page(doc)
    add_contents_page(doc)

    doc.add_heading("1. 方案目标", level=1)
    doc.add_paragraph("为保障系统长期稳定、安全、高效运行，建立覆盖基础设施、应用服务、数据库、中间件、网络安全、监控告警、应急响应、变更发布、备份灾备和持续优化的运维保障体系，确保生产系统整体可用性达到 99.99% SLA 要求。")
    add_bullets(doc, [
        "稳定性目标：核心业务链路持续可用，故障可快速发现、快速止血、快速恢复。",
        "安全性目标：生产访问最小权限、关键操作可审计、敏感数据可追踪、异常行为可拦截。",
        "连续性目标：关键数据具备备份、恢复、主备切换和灾备演练能力。",
        "可运营目标：指标、流程、角色、预案和复盘机制可落地、可考核、可持续优化。",
    ])

    doc.add_heading("2. 适用范围", level=1)
    add_kv_table(doc, [
        ("系统范围", "核心业务系统、管理后台、开放接口、定时任务、批处理任务、LLM/ASR/TTS 能力服务及其依赖服务"),
        ("基础设施", "IDC 物理机、GPU 服务器、私有化虚拟化平台、负载均衡、DNS、WAF、网络链路、共享存储、对象存储"),
        ("数据与中间件", "关系型数据库、缓存、消息队列、搜索服务、向量检索、配置中心、文件存储"),
        ("运维流程", "监控告警、发布变更、模型版本管理、故障响应、巡检、容量管理、备份恢复、安全审计"),
    ])

    doc.add_heading("3. SLA 指标体系", level=1)
    add_table(doc, ["指标项", "目标值", "说明"], [
        ("系统可用性", "≥ 99.99%", "按月统计，计划维护窗口需提前确认并公告"),
        ("月度最大不可用时间", "≤ 4.38 分钟", "99.99% 对应每月约 4.38 分钟不可用预算"),
        ("年度最大不可用时间", "≤ 52.56 分钟", "年度维度用于趋势复盘和管理评审"),
        ("P1 响应时间", "≤ 5 分钟", "多通道通知值班人员并确认接单"),
        ("P1 RTO", "≤ 15 分钟", "核心链路建议目标 ≤ 5 分钟"),
        ("关键数据 RPO", "≤ 5 分钟", "核心交易类数据建议接近 0"),
        ("告警发现时间", "≤ 1 分钟", "关键指标实时采集和高频拨测"),
        ("重大变更回退时间", "≤ 10 分钟", "发布前必须验证回滚路径"),
    ], [1.45, 1.35, 3.57])

    add_topology(doc)

    doc.add_heading("5. 高可用保障措施", level=1)
    doc.add_heading("5.1 基础设施高可用", level=2)
    add_bullets(doc, [
        "生产资源部署在两个及以上可用区，核心节点避免单实例运行。",
        "公网入口、负载均衡、NAT、专线或 VPN 按业务重要性配置冗余链路。",
        "关键云资源启用健康检查、自动恢复、自动扩缩容和容量保护策略。",
        "基础设施变更纳入变更审批，保留网络、安全组、路由和证书变更记录。",
    ])
    doc.add_heading("5.2 应用高可用", level=2)
    add_bullets(doc, [
        "应用无状态化，Session、缓存、文件、任务状态外置到高可用存储。",
        "应用至少多副本部署，跨可用区承载流量，单节点故障不影响整体服务。",
        "核心接口配置超时、重试、限流、熔断、隔离和降级策略。",
        "关键任务具备幂等能力和补偿机制，避免重复消费或重复写入。",
    ])
    doc.add_heading("5.3 数据库与中间件高可用", level=2)
    add_bullets(doc, [
        "数据库采用主备、集群或云托管高可用架构，启用自动或半自动故障转移。",
        "Redis、消息队列、搜索服务采用集群模式，关键数据开启持久化或多副本。",
        "消息队列配置死信队列、重试队列、消费延迟监控和积压告警。",
        "数据库慢 SQL、锁等待、复制延迟、连接池耗尽等指标纳入实时监控。",
    ])
    doc.add_heading("5.4 LLM / ASR / TTS 高可用", level=2)
    add_bullets(doc, [
        "模型服务统一经模型网关接入，调用方不直接绑定单个模型实例或单台推理节点。",
        "LLM、ASR、TTS 分别配置主备模型、备用节点和降级策略，支持按业务优先级限流。",
        "对长文本、长音频、批量合成等重负载请求设置队列、超时、最大长度和并发上限。",
        "Prompt、模型版本、音色模型、ASR 后处理规则均纳入版本管理，支持快速回退。",
        "关键智能能力必须具备业务兜底方案，例如模板回复、缓存音频、人工处理或延后处理。",
    ])
    doc.add_heading("5.5 IDC 与私有化虚拟化高可用", level=2)
    add_bullets(doc, [
        "物理机按机柜、交换机、存储链路和电源故障域规划，核心业务跨故障域分布。",
        "私有化虚拟化平台启用集群 HA、虚机反亲和、资源预留和管理面备份。",
        "共享存储、核心交换、负载均衡、防火墙和出口链路具备冗余路径，定期验证切换有效性。",
        "GPU 推理资源池与普通业务资源池隔离，避免资源争抢导致核心业务和模型服务同时退化。",
    ])

    doc.add_heading("6. 监控告警体系", level=1)
    add_table(doc, ["监控域", "关键指标", "告警目标"], [
        ("基础资源", "CPU、内存、磁盘、网络、负载、进程、容器状态", "提前发现容量和节点风险"),
        ("应用服务", "QPS、响应时间、错误率、接口状态、线程池、连接池", "定位服务质量下降和异常接口"),
        ("数据库", "连接数、慢查询、锁等待、复制延迟、TPS、缓存命中率", "防止数据库成为核心瓶颈"),
        ("中间件", "Redis 命中率、MQ 积压、消费延迟、节点状态", "保障异步链路和缓存链路稳定"),
        ("业务指标", "登录成功率、下单成功率、支付成功率、任务成功率", "用业务结果校验系统真实可用性"),
        ("安全指标", "异常登录、越权、WAF 拦截、暴力破解、异常流量", "快速识别攻击和滥用行为"),
    ], [1.15, 3.1, 2.12])
    add_bullets(doc, [
        "P1/P2 告警通过电话、短信、IM、值班系统多通道通知。",
        "告警规则必须配置责任人、升级路径、恢复通知和抑制策略。",
        "关键页面、核心接口和第三方依赖配置端到端拨测。",
        "告警恢复后记录处理动作，未闭环问题进入稳定性问题台账。",
    ])
    add_specialized_monitoring(doc)

    doc.add_heading("7. 变更与发布管理", level=1)
    add_numbered(doc, [
        "提交变更申请，明确变更内容、影响范围、风险等级、执行窗口和负责人。",
        "完成代码审核、测试验证、配置核对和数据库脚本审查。",
        "制定发布计划、灰度策略、监控观察项和回滚方案。",
        "低峰期执行灰度或滚动发布，严禁高风险变更直接全量上线。",
        "发布后观察核心业务指标、错误率、延迟、资源水位和用户反馈。",
        "发布完成后记录结果；异常发布必须复盘并修正发布流程。",
    ])

    doc.add_heading("8. 备份与灾备", level=1)
    add_table(doc, ["对象", "备份策略", "恢复要求"], [
        ("数据库", "每日全量备份，实时或准实时增量备份，跨可用区保存", "定期恢复演练，关键数据 RPO ≤ 5 分钟"),
        ("配置文件", "每次变更前后自动留存版本", "支持分钟级回退到历史版本"),
        ("应用制品", "保留最近多个稳定版本和构建元数据", "支持快速回滚到上一稳定版本"),
        ("模型与规则", "模型文件、Prompt 模板、ASR/TTS 配置、音色模型、后处理规则保留版本", "支持快速回退到上一稳定模型或规则版本"),
        ("虚拟化平台", "管理面配置、虚机模板、关键虚机快照、网络与存储配置定期备份", "管理面故障后可恢复，关键虚机具备重建依据"),
        ("对象文件", "开启多副本、版本控制和生命周期策略", "误删可恢复，归档可追踪"),
        ("日志数据", "按合规周期集中归档", "故障复盘和安全审计可检索"),
    ], [1.05, 3.35, 1.97])
    add_bullets(doc, [
        "每季度至少开展一次数据库恢复演练和灾备切换演练。",
        "灾备预案应脚本化、清单化，明确前置条件、执行步骤、验证标准和回切流程。",
        "备份成功不等于可恢复，必须用恢复演练验证备份有效性。",
    ])

    doc.add_heading("9. 安全保障", level=1)
    add_bullets(doc, [
        "生产环境采用最小权限模型，运维入口统一通过堡垒机或受控通道访问。",
        "管理后台、云控制台、代码仓库、CI/CD 平台启用 MFA。",
        "敏感配置通过密钥管理系统保存，禁止明文写入代码仓库和镜像。",
        "对外服务接入 WAF、防 DDoS、访问频率限制和异常行为识别。",
        "定期开展漏洞扫描、基线检查、依赖风险治理和权限复核。",
        "生产操作、数据导出、权限变更和安全事件必须保留审计记录。",
    ])

    doc.add_heading("10. 容量与性能保障", level=1)
    add_bullets(doc, [
        "核心资源平均水位建议控制在 70% 以下，高峰前保留不低于 30% 冗余。",
        "重大活动、版本发布、营销推广前完成压测和容量评估。",
        "对慢接口、慢 SQL、队列积压、缓存穿透和热点 Key 建立专项优化机制。",
        "每月输出容量趋势报告，包含未来 1 至 3 个月扩容建议。",
        "关键服务具备自动扩容策略，并设置最大扩容边界避免成本失控。",
    ])

    add_incident_playbooks(doc)

    doc.add_heading("12. 巡检制度", level=1)
    add_table(doc, ["频率", "巡检内容", "输出"], [
        ("每日", "核心服务状态、LLM/ASR/TTS 失败率与延迟、告警、备份结果、资源水位、业务成功率、安全事件", "日巡检记录和异常处理单"),
        ("每周", "容量趋势、GPU/虚拟化资源池、慢 SQL、发布记录、中间件健康、告警有效性", "周报和问题台账更新"),
        ("每月", "SLA 达成、故障统计、复盘闭环、安全基线、灾备能力、IDC 硬件风险、成本", "月度运维报告"),
        ("每季度", "灾备演练、虚拟化 HA 演练、模型降级演练、权限复核、重大风险评审、容量规划", "季度稳定性评审报告"),
    ], [0.8, 4.05, 1.52])

    doc.add_heading("13. 运维组织与值班机制", level=1)
    add_table(doc, ["角色", "职责"], [
        ("应急负责人", "统一指挥，协调资源，决策恢复方案，对外同步故障状态"),
        ("一线值班", "接收告警、初步判断、执行标准化止血动作、拉起应急群"),
        ("运维负责人", "处理基础设施、网络、容器、发布回滚和资源扩容"),
        ("开发负责人", "定位应用代码、接口、任务、依赖调用和业务逻辑异常"),
        ("算法/模型负责人", "处理 LLM、ASR、TTS 模型效果、模型版本、Prompt、音色和推理质量异常"),
        ("基础设施负责人", "处理 IDC 物理机、GPU 服务器、虚拟化平台、共享存储和机房网络故障"),
        ("数据库负责人", "处理数据库性能、主备切换、备份恢复和数据一致性"),
        ("安全负责人", "判断攻击、入侵、权限异常和安全加固动作"),
        ("业务负责人", "评估业务影响，协调用户通知、业务补偿和客户沟通"),
    ], [1.25, 5.12])
    add_bullets(doc, [
        "建立 7×24 小时值班机制，明确一线、二线专家和管理升级路径。",
        "交接班必须同步未关闭问题、风险变更、重点观察项和待升级事项。",
        "节假日、重大活动和高风险变更期间启动专项保障值守。",
    ])

    doc.add_heading("14. 故障复盘与持续改进", level=1)
    add_bullets(doc, [
        "所有 P1/P2 故障必须在规定时间内完成复盘，输出时间线、影响范围、根因、恢复过程、改进项和责任人。",
        "复盘以系统性改进为目标，不以追责为主要目的。",
        "重复故障必须专项治理，必要时纳入架构改造、自动化恢复或流程调整。",
        "每月进行 SLA 统计，评估可用性、故障次数、平均响应时间、平均恢复时间、变更成功率和告警准确率。",
    ])

    doc.add_heading("15. 附录 A：SLA 统计口径", level=1)
    add_callout(doc, "计算公式", "系统可用性 = （统计周期总时间 - 不可用时间）/ 统计周期总时间 × 100%。不可用时间原则上以核心业务链路不可用、关键接口错误率超过阈值且影响用户成功完成业务为准。")
    add_bullets(doc, [
        "计划维护窗口需提前通知并经业务确认，可按合同或制度约定从不可用时间中剔除。",
        "第三方依赖导致的不可用是否计入 SLA，应在服务边界或合同中明确。",
        "部分功能不可用时，应按影响范围、用户比例和业务重要性折算或分级统计。",
    ])

    doc.add_heading("16. 附录 B：上线前稳定性检查清单", level=1)
    add_table(doc, ["检查项", "是否必需", "验收标准"], [
        ("多副本部署", "是", "任一实例下线不影响整体服务"),
        ("健康检查", "是", "异常实例可自动摘除并恢复"),
        ("回滚方案", "是", "10 分钟内可回退到上一稳定版本"),
        ("核心监控", "是", "覆盖错误率、延迟、QPS、资源水位和业务成功率"),
        ("模型监控", "是", "覆盖 LLM/ASR/TTS 失败率、延迟、队列、GPU、质量抽检和降级状态"),
        ("基础设施监控", "是", "覆盖物理机、GPU、虚拟化集群、共享存储、网络、电力和制冷"),
        ("告警责任人", "是", "P1/P2 告警可触达值班人员"),
        ("数据库备份", "是", "备份任务成功且恢复演练可用"),
        ("模型与配置备份", "是", "模型文件、Prompt、音色、ASR/TTS 规则可回退"),
        ("压测结果", "按需", "峰值容量满足业务预测并保留冗余"),
        ("安全检查", "是", "无高危漏洞，敏感配置未明文暴露"),
    ], [2.0, 1.0, 3.37])

    doc.add_heading("17. 附录 C：故障复盘模板", level=1)
    add_kv_table(doc, [
        ("故障编号", "按年月日和序号生成，例如 INC-20260520-001"),
        ("故障等级", "P1 / P2 / P3 / P4"),
        ("开始时间 / 恢复时间", "记录首次告警、确认、止血、恢复和关闭时间"),
        ("影响范围", "受影响系统、用户比例、核心业务影响、数据影响"),
        ("智能能力影响", "LLM、ASR、TTS 是否受影响，涉及模型版本、Prompt、音色或音频链路"),
        ("基础设施影响", "涉及物理机、GPU、宿主机、虚拟化集群、共享存储、网络或机房环境"),
        ("直接原因", "触发故障的直接技术原因"),
        ("根本原因", "架构、流程、监控、容量、代码或管理层面的深层原因"),
        ("恢复动作", "切流、扩容、回滚、降级、修复、数据恢复等动作"),
        ("改进项", "明确责任人、完成时间和验证方式"),
    ])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
